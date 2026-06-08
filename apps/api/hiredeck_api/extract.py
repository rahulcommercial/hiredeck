"""Text extraction from uploaded files (PDF / DOCX / TXT)."""

from __future__ import annotations

import io

import docx  # type: ignore[import-untyped]
from pypdf import PdfReader


def extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — pypdf occasionally trips on malformed pages
            continue
    return "\n\n".join(p.strip() for p in parts if p.strip())


def extract_docx(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # also pull table cells, which many resumes use for layout
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_text(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_pdf(data)
    if lower.endswith(".docx"):
        return extract_docx(data)
    if lower.endswith(".txt"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {filename}")
